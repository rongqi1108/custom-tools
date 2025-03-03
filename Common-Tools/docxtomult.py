import os
import re
import tempfile
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.image.exceptions import UnrecognizedImageError


# 函数：去掉文件名中的非法字符
def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '_', filename)


# 函数：复制图片
def copy_images(source_doc, target_doc):
    for rel in source_doc.part.rels.values():
        if "image" in rel.target_ref:
            # 获取图片的二进制数据
            image_data = rel.target_part.blob

            # 使用临时文件保存图片
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                temp_file.write(image_data)
                temp_file_path = temp_file.name

            try:
                # 尝试插入图片到目标文档，指定默认宽度和高度
                target_doc.add_picture(temp_file_path, width=Inches(4), height=Inches(3))
            except UnrecognizedImageError:
                # 捕获图片无法识别的错误并跳过
                print(f"警告: 无法识别图片 {temp_file_path}，跳过此图片。")
            except ZeroDivisionError:
                # 捕获 DPI 缺失的错误并使用默认尺寸
                print(f"警告: 图片 {temp_file_path} 的 DPI 信息丢失，插入时使用默认尺寸。")
                target_doc.add_picture(temp_file_path, width=Inches(4), height=Inches(3))

            # 删除临时文件
            os.remove(temp_file_path)


# 函数：复制表格
def copy_table(source_table, target_doc):
    if target_doc is None:
        print("目标文档未初始化!")
        return
    table = target_doc.add_table(rows=0, cols=len(source_table.columns))
    for row in source_table.rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row.cells):
            row_cells[i].text = cell.text
    return table


def split_docx_by_headings(input_path):
    # 检查文件是否存在
    if not os.path.exists(input_path):
        print(f"文件路径不存在: {input_path}")
        return

    # 打开输入文档
    doc = Document(input_path)

    # 获取输入文档的目录
    input_dir = os.path.dirname(input_path)

    # 初始化
    paragraphs = doc.paragraphs
    current_doc = None
    current_title = None

    # 遍历所有段落，查找标题并拆分文档
    for para in paragraphs:
        # 假设标题是使用 "Heading" 样式的（如：Heading 1, Heading 2）
        if para.style.name.startswith('Heading'):
            # 如果有当前文档正在处理中，保存当前文档
            if current_doc:
                sanitized_title = sanitize_filename(current_title)
                output_file = os.path.join(input_dir, f"{sanitized_title}.docx")
                current_doc.save(output_file)
                print(f"保存文档: {output_file}")

            # 新建一个空文档，并设置标题为当前段落
            current_title = para.text.strip()  # 获取标题文本
            current_doc = Document()
            current_doc.add_paragraph(para.text)  # 将标题添加到新文档

        else:
            # 否则，将段落内容加入当前文档
            if current_doc:
                current_doc.add_paragraph(para.text)

        # 复制表格和图片 (在每个标题之后)
        if current_doc:  # 只有在current_doc已经初始化后才复制表格和图片
            for table in doc.tables:
                copy_table(table, current_doc)
            copy_images(doc, current_doc)

    # 保存最后一个文档
    if current_doc:
        sanitized_title = sanitize_filename(current_title)
        output_file = os.path.join(input_dir, f"{sanitized_title}.docx")
        current_doc.save(output_file)
        print(f"保存文档: {output_file}")

    print("文档拆分完成！")


# 使用函数拆分文档
input_path = r'F:\\cpp\\面经\\split\\面经.docx'  # 替换为你自己的文件路径
split_docx_by_headings(input_path)

